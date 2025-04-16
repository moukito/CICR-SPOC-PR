"""
Processor for DOCX documents.

This module provides functionality to extract content from Microsoft Word (DOCX) files
and convert it to Document objects for indexing and querying within the
document processing system.

The DocxProcessor class extracts both paragraph text and tabular data from Word documents.
It preserves document structure and formats tables in a readable way with
appropriate table identifiers to maintain document context.

Classes:
    DocxProcessor: Processes DOCX files into Document objects.

Methods:
    DocxProcessor.process_file: Extracts content from DOCX files and creates
                               Document objects with appropriate metadata.
    DocxProcessor.extract_docx_text: Handles the extraction of text and tables
                                    from Word documents.

Dependencies:
    python-docx: Used for DOCX content extraction.

This processor is part of the document processing subsystem and is automatically
selected for DOCX files by the get_document_processor factory function.
"""

import docx
from llama_index.core import Document


class DocxProcessor:
    """
    Processor for handling DOCX (Microsoft Word) files extraction.
    Converts DOCX content to Document objects with appropriate metadata,
    preserving both textual content and tabular data structure.
    """

    def process_file(self, file_path):
        """
        Extracts content from a DOCX file.

        Args:
            file_path: Path to the DOCX file

        Returns:
            Document: Document containing text and metadata
        """
        content = self.extract_docx_text_and_tables(file_path)
        return Document(text=content, metadata={"filename": file_path, "type": "docx"})

    @staticmethod
    def extract_docx_text_and_tables(file_path):
        """
        Extracts text and tables from a DOCX file.

        This method processes all paragraphs and tables in the document,
        preserving the structure. Tables are properly labeled with identifying
        information to maintain document organization.

        Args:
            file_path: Path to the DOCX file

        Returns:
            str: Formatted string containing all extracted text and tables
        """
        doc = docx.Document(file_path)
        full_text = []

        # Extract text from paragraphs
        for para in doc.paragraphs:
            full_text.append(para.text)

        # Extract text from tables
        for i, table in enumerate(doc.tables):
            full_text.append(f"\n[Tableau {i + 1}]\n")
            for row in table.rows:
                row_content = []
                for cell in row.cells:
                    row_content.append(cell.text)
                full_text.append(" | ".join(row_content))

        return "\n".join(full_text)
